"""
Document Generator Service.

Generates professional documents in multiple formats:
- PDF (for auditors, official reports)
- Excel (for accountants, data analysis)
- Word (for editing, contracts)
- Markdown (for developers)

Designed for finance/operations teams and auditors.
"""
import io
import json
import logging
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, List, Optional, Union
from enum import Enum

logger = logging.getLogger(__name__)


class DocumentFormat(str, Enum):
    """Supported document formats."""
    PDF = "pdf"
    EXCEL = "xlsx"
    WORD = "docx"
    MARKDOWN = "md"
    JSON = "json"
    CSV = "csv"


@dataclass
class DocumentConfig:
    """Document generation configuration."""
    format: DocumentFormat
    language: str = "en"
    include_metrics: bool = True
    include_recommendations: bool = True
    include_cost_breakdown: bool = True
    template: str = "standard"
    company_name: str = ""
    logo_path: Optional[str] = None


class DocumentGenerator:
    """
    Generates professional documents from analysis results.

    Usage:
        generator = DocumentGenerator()
        pdf_bytes = generator.generate_pdf(analysis_result, config)
        excel_bytes = generator.generate_excel(analysis_result, config)
    """

    def __init__(self):
        self._check_dependencies()

    def _check_dependencies(self) -> Dict[str, bool]:
        """Check available dependencies."""
        available = {}

        try:
            import reportlab
            available["pdf"] = True
        except ImportError:
            available["pdf"] = False

        try:
            import openpyxl
            available["excel"] = True
        except ImportError:
            available["excel"] = False

        try:
            from docx import Document
            available["word"] = True
        except ImportError:
            available["word"] = False

        return available

    def generate(
        self,
        data: Dict[str, Any],
        format: DocumentFormat,
        config: Optional[DocumentConfig] = None,
    ) -> bytes:
        """Generate document in specified format."""
        config = config or DocumentConfig(format=format)

        if format == DocumentFormat.PDF:
            return self.generate_pdf(data, config)
        elif format == DocumentFormat.EXCEL:
            return self.generate_excel(data, config)
        elif format == DocumentFormat.WORD:
            return self.generate_word(data, config)
        elif format == DocumentFormat.MARKDOWN:
            return self.generate_markdown(data, config).encode()
        elif format == DocumentFormat.JSON:
            return json.dumps(data, indent=2, default=str).encode()
        elif format == DocumentFormat.CSV:
            return self.generate_csv(data, config).encode()
        else:
            raise ValueError(f"Unsupported format: {format}")

    def generate_pdf(self, data: Dict[str, Any], config: DocumentConfig) -> bytes:
        """Generate PDF report."""
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import cm
            from reportlab.platypus import (
                SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
                PageBreak, Image
            )
            from reportlab.lib.enums import TA_CENTER, TA_LEFT
        except ImportError:
            logger.warning("reportlab not installed, falling back to text-based PDF")
            return self._generate_text_pdf(data, config)

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            rightMargin=2*cm,
            leftMargin=2*cm,
            topMargin=2*cm,
            bottomMargin=2*cm,
        )

        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(
            name='Title2',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30,
            alignment=TA_CENTER,
        ))
        styles.add(ParagraphStyle(
            name='Subtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.grey,
            alignment=TA_CENTER,
        ))

        elements = []

        # Title
        elements.append(Paragraph("Repository Audit Report", styles['Title2']))
        elements.append(Paragraph(
            f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}",
            styles['Subtitle']
        ))
        elements.append(Spacer(1, 30))

        # Executive Summary
        elements.append(Paragraph("Executive Summary", styles['Heading2']))
        elements.append(Spacer(1, 10))

        summary_data = [
            ["Metric", "Value"],
            ["Product Level", data.get("product_level", "N/A")],
            ["Complexity", data.get("complexity", "N/A")],
            ["Repo Health", f"{data.get('repo_health', {}).get('total', 0)}/12"],
            ["Tech Debt", f"{data.get('tech_debt', {}).get('total', 0)}/15"],
        ]

        summary_table = Table(summary_data, colWidths=[8*cm, 6*cm])
        summary_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1e40af')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 11),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.HexColor('#f8fafc')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#e2e8f0')),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1), [colors.white, colors.HexColor('#f8fafc')]),
        ]))
        elements.append(summary_table)
        elements.append(Spacer(1, 20))

        # Repo Health Details
        elements.append(Paragraph("Repo Health Breakdown", styles['Heading2']))
        elements.append(Spacer(1, 10))

        health = data.get("repo_health", {})
        health_data = [
            ["Category", "Score", "Max"],
            ["Documentation", str(health.get("documentation", 0)), "3"],
            ["Structure", str(health.get("structure", 0)), "3"],
            ["Runability", str(health.get("runability", 0)), "3"],
            ["History", str(health.get("history", 0)), "3"],
            ["Total", str(health.get("total", 0)), "12"],
        ]

        health_table = Table(health_data, colWidths=[6*cm, 4*cm, 4*cm])
        health_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#059669')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#ecfdf5')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#d1fae5')),
        ]))
        elements.append(health_table)
        elements.append(Spacer(1, 20))

        # Tech Debt Details
        elements.append(Paragraph("Tech Debt Breakdown", styles['Heading2']))
        elements.append(Spacer(1, 10))

        debt = data.get("tech_debt", {})
        debt_data = [
            ["Category", "Score", "Max"],
            ["Architecture", str(debt.get("architecture", 0)), "3"],
            ["Code Quality", str(debt.get("code_quality", 0)), "3"],
            ["Testing", str(debt.get("testing", 0)), "3"],
            ["Infrastructure", str(debt.get("infrastructure", 0)), "3"],
            ["Security", str(debt.get("security", 0)), "3"],
            ["Total", str(debt.get("total", 0)), "15"],
        ]

        debt_table = Table(debt_data, colWidths=[6*cm, 4*cm, 4*cm])
        debt_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#dc2626')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
            ('BACKGROUND', (0, -1), (-1, -1), colors.HexColor('#fef2f2')),
            ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#fecaca')),
        ]))
        elements.append(debt_table)
        elements.append(Spacer(1, 20))

        # Cost Estimation
        if config.include_cost_breakdown:
            elements.append(Paragraph("Cost Estimation", styles['Heading2']))
            elements.append(Spacer(1, 10))

            cost = data.get("cost", {})
            cost_data = [
                ["Metric", "Value"],
                ["Estimated Hours", str(cost.get("hours", "N/A"))],
                ["Cost Range", cost.get("formatted", "N/A")],
                ["Profile", cost.get("profile", "N/A")],
            ]

            cost_table = Table(cost_data, colWidths=[8*cm, 6*cm])
            cost_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#7c3aed')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#ddd6fe')),
            ]))
            elements.append(cost_table)
            elements.append(Spacer(1, 20))

        # Recommendations
        if config.include_recommendations:
            recs = data.get("recommendations", [])
            if recs:
                elements.append(Paragraph("Recommendations", styles['Heading2']))
                elements.append(Spacer(1, 10))
                for rec in recs:
                    elements.append(Paragraph(f"• {rec}", styles['Normal']))
                    elements.append(Spacer(1, 5))

        # Footer
        elements.append(Spacer(1, 30))
        elements.append(Paragraph(
            f"Generated by Repo Auditor | {config.company_name or 'https://repo-auditor.com'}",
            styles['Subtitle']
        ))

        doc.build(elements)
        return buffer.getvalue()

    def _generate_text_pdf(self, data: Dict[str, Any], config: DocumentConfig) -> bytes:
        """Fallback text-based PDF using markdown."""
        md = self.generate_markdown(data, config)
        return f"PDF generation requires reportlab. Install with: pip install reportlab\n\n{md}".encode()

    def generate_excel(self, data: Dict[str, Any], config: DocumentConfig) -> bytes:
        """Generate Excel workbook."""
        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, Fill, PatternFill, Alignment, Border, Side
            from openpyxl.utils import get_column_letter
        except ImportError:
            logger.warning("openpyxl not installed")
            return self.generate_csv(data, config).encode()

        wb = Workbook()

        # Summary sheet
        ws = wb.active
        ws.title = "Summary"

        # Styles
        header_fill = PatternFill(start_color="1e40af", end_color="1e40af", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        border = Border(
            left=Side(style='thin'),
            right=Side(style='thin'),
            top=Side(style='thin'),
            bottom=Side(style='thin')
        )

        # Title
        ws['A1'] = "Repository Audit Report"
        ws['A1'].font = Font(size=18, bold=True)
        ws.merge_cells('A1:D1')

        ws['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ws['A3'] = f"Repository: {data.get('repo_path', 'N/A')}"

        # Summary table
        row = 5
        headers = ["Metric", "Value", "Max", "Percentage"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border

        # Data rows
        health = data.get("repo_health", {})
        debt = data.get("tech_debt", {})

        summary_rows = [
            ("Product Level", data.get("product_level", "N/A"), "-", "-"),
            ("Complexity", data.get("complexity", "N/A"), "-", "-"),
            ("Repo Health", health.get("total", 0), 12, f"{round(health.get('total', 0)/12*100)}%"),
            ("Tech Debt", debt.get("total", 0), 15, f"{round(debt.get('total', 0)/15*100)}%"),
            ("", "", "", ""),
            ("Documentation", health.get("documentation", 0), 3, ""),
            ("Structure", health.get("structure", 0), 3, ""),
            ("Runability", health.get("runability", 0), 3, ""),
            ("History", health.get("history", 0), 3, ""),
            ("", "", "", ""),
            ("Architecture", debt.get("architecture", 0), 3, ""),
            ("Code Quality", debt.get("code_quality", 0), 3, ""),
            ("Testing", debt.get("testing", 0), 3, ""),
            ("Infrastructure", debt.get("infrastructure", 0), 3, ""),
            ("Security", debt.get("security", 0), 3, ""),
        ]

        for i, (metric, value, max_val, pct) in enumerate(summary_rows, row + 1):
            ws.cell(row=i, column=1, value=metric).border = border
            ws.cell(row=i, column=2, value=value).border = border
            ws.cell(row=i, column=3, value=max_val).border = border
            ws.cell(row=i, column=4, value=pct).border = border

        # Cost sheet
        ws_cost = wb.create_sheet("Cost Estimation")
        cost = data.get("cost", {})

        ws_cost['A1'] = "Cost Estimation"
        ws_cost['A1'].font = Font(size=14, bold=True)

        cost_data = [
            ("Hours", cost.get("hours", 0)),
            ("Currency", cost.get("currency", "EUR")),
            ("Min Cost", cost.get("min", 0)),
            ("Max Cost", cost.get("max", 0)),
            ("Profile", cost.get("profile", "Standard")),
        ]

        for i, (label, value) in enumerate(cost_data, 3):
            ws_cost.cell(row=i, column=1, value=label)
            ws_cost.cell(row=i, column=2, value=value)

        # Code Metrics sheet
        if "code" in data:
            ws_code = wb.create_sheet("Code Metrics")
            ws_code['A1'] = "Code Metrics by Language"
            ws_code['A1'].font = Font(size=14, bold=True)

            ws_code['A3'] = "Language"
            ws_code['B3'] = "Files"
            ws_code['C3'] = "LOC"
            for cell in [ws_code['A3'], ws_code['B3'], ws_code['C3']]:
                cell.fill = header_fill
                cell.font = header_font

            row = 4
            for lang, info in data["code"].get("files", {}).items():
                ws_code.cell(row=row, column=1, value=lang)
                ws_code.cell(row=row, column=2, value=info.get("count", 0))
                ws_code.cell(row=row, column=3, value=info.get("loc", 0))
                row += 1

        # Adjust column widths
        for ws in wb.worksheets:
            for column in ws.columns:
                max_length = 0
                column_letter = get_column_letter(column[0].column)
                for cell in column:
                    if cell.value:
                        max_length = max(max_length, len(str(cell.value)))
                ws.column_dimensions[column_letter].width = max(max_length + 2, 12)

        buffer = io.BytesIO()
        wb.save(buffer)
        return buffer.getvalue()

    def generate_word(self, data: Dict[str, Any], config: DocumentConfig) -> bytes:
        """Generate Word document."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.table import WD_TABLE_ALIGNMENT
        except ImportError:
            logger.warning("python-docx not installed")
            return self.generate_markdown(data, config).encode()

        doc = Document()

        # Title
        title = doc.add_heading('Repository Audit Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        health = data.get("repo_health", {})
        debt = data.get("tech_debt", {})

        rows_data = [
            ("Product Level", data.get("product_level", "N/A")),
            ("Complexity", data.get("complexity", "N/A")),
            ("Repo Health", f"{health.get('total', 0)}/12 ({round(health.get('total', 0)/12*100)}%)"),
            ("Tech Debt", f"{debt.get('total', 0)}/15 ({round(debt.get('total', 0)/15*100)}%)"),
            ("Cost Estimate", data.get("cost", {}).get("formatted", "N/A")),
        ]

        for i, (label, value) in enumerate(rows_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

        # Repo Health
        doc.add_heading('Repo Health Breakdown', level=1)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        health_rows = [
            ("Category", "Score", "Max"),
            ("Documentation", str(health.get("documentation", 0)), "3"),
            ("Structure", str(health.get("structure", 0)), "3"),
            ("Runability", str(health.get("runability", 0)), "3"),
            ("History", str(health.get("history", 0)), "3"),
        ]

        for i, row in enumerate(health_rows):
            for j, val in enumerate(row):
                table.rows[i].cells[j].text = val

        doc.add_paragraph()

        # Tech Debt
        doc.add_heading('Tech Debt Breakdown', level=1)

        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        debt_rows = [
            ("Category", "Score", "Max"),
            ("Architecture", str(debt.get("architecture", 0)), "3"),
            ("Code Quality", str(debt.get("code_quality", 0)), "3"),
            ("Testing", str(debt.get("testing", 0)), "3"),
            ("Infrastructure", str(debt.get("infrastructure", 0)), "3"),
            ("Security", str(debt.get("security", 0)), "3"),
        ]

        for i, row in enumerate(debt_rows):
            for j, val in enumerate(row):
                table.rows[i].cells[j].text = val

        doc.add_paragraph()

        # Cost Estimation
        if config.include_cost_breakdown:
            doc.add_heading('Cost Estimation', level=1)
            cost = data.get("cost", {})

            p = doc.add_paragraph()
            p.add_run(f"Estimated Hours: ").bold = True
            p.add_run(f"{cost.get('hours', 'N/A')}")

            p = doc.add_paragraph()
            p.add_run(f"Cost Range: ").bold = True
            p.add_run(f"{cost.get('formatted', 'N/A')}")

            p = doc.add_paragraph()
            p.add_run(f"Pricing Profile: ").bold = True
            p.add_run(f"{cost.get('profile', 'N/A')}")

        doc.add_paragraph()

        # Recommendations
        if config.include_recommendations:
            recs = data.get("recommendations", [])
            if recs:
                doc.add_heading('Recommendations', level=1)
                for rec in recs:
                    doc.add_paragraph(rec, style='List Bullet')

        # Footer
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated by Repo Auditor | {config.company_name or 'repo-auditor.com'}")
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(9)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def generate_markdown(self, data: Dict[str, Any], config: DocumentConfig) -> str:
        """Generate Markdown report."""
        health = data.get("repo_health", {})
        debt = data.get("tech_debt", {})
        cost = data.get("cost", {})
        code = data.get("code", {})

        return f"""# Repository Audit Report

**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}
**Repository:** {data.get('repo_path', 'N/A')}

---

## Executive Summary

| Metric | Value |
|--------|-------|
| Product Level | {data.get('product_level', 'N/A')} |
| Complexity | {data.get('complexity', 'N/A')} |
| Repo Health | {health.get('total', 0)}/12 ({round(health.get('total', 0)/12*100)}%) |
| Tech Debt | {debt.get('total', 0)}/15 ({round(debt.get('total', 0)/15*100)}%) |

---

## Repo Health

| Category | Score |
|----------|-------|
| Documentation | {health.get('documentation', 0)}/3 |
| Structure | {health.get('structure', 0)}/3 |
| Runability | {health.get('runability', 0)}/3 |
| History | {health.get('history', 0)}/3 |
| **Total** | **{health.get('total', 0)}/12** |

---

## Tech Debt

| Category | Score |
|----------|-------|
| Architecture | {debt.get('architecture', 0)}/3 |
| Code Quality | {debt.get('code_quality', 0)}/3 |
| Testing | {debt.get('testing', 0)}/3 |
| Infrastructure | {debt.get('infrastructure', 0)}/3 |
| Security | {debt.get('security', 0)}/3 |
| **Total** | **{debt.get('total', 0)}/15** |

---

## Cost Estimation

| Metric | Value |
|--------|-------|
| Hours | {cost.get('hours', 'N/A')} |
| Cost Range | {cost.get('formatted', 'N/A')} |
| Profile | {cost.get('profile', 'N/A')} |

---

## Code Metrics

| Metric | Value |
|--------|-------|
| Total Files | {code.get('total_files', 0)} |
| Total LOC | {code.get('total_loc', 0):,} |
| Test Files | {code.get('test_files', 0)} |

---

## Recommendations

{chr(10).join(f'- {r}' for r in data.get('recommendations', ['No critical issues.']))}

---

*Generated by Repo Auditor*
"""

    def generate_csv(self, data: Dict[str, Any], config: DocumentConfig) -> str:
        """Generate CSV for Excel import."""
        health = data.get("repo_health", {})
        debt = data.get("tech_debt", {})
        cost = data.get("cost", {})

        lines = [
            "Category,Metric,Value,Max,Notes",
            f"Summary,Product Level,{data.get('product_level', '')},,-",
            f"Summary,Complexity,{data.get('complexity', '')},,-",
            f"Summary,Repo Health,{health.get('total', 0)},12,{round(health.get('total', 0)/12*100)}%",
            f"Summary,Tech Debt,{debt.get('total', 0)},15,{round(debt.get('total', 0)/15*100)}%",
            "",
            f"Repo Health,Documentation,{health.get('documentation', 0)},3,",
            f"Repo Health,Structure,{health.get('structure', 0)},3,",
            f"Repo Health,Runability,{health.get('runability', 0)},3,",
            f"Repo Health,History,{health.get('history', 0)},3,",
            "",
            f"Tech Debt,Architecture,{debt.get('architecture', 0)},3,",
            f"Tech Debt,Code Quality,{debt.get('code_quality', 0)},3,",
            f"Tech Debt,Testing,{debt.get('testing', 0)},3,",
            f"Tech Debt,Infrastructure,{debt.get('infrastructure', 0)},3,",
            f"Tech Debt,Security,{debt.get('security', 0)},3,",
            "",
            f"Cost,Hours,{cost.get('hours', '')},,",
            f"Cost,Min,{cost.get('min', '')},{cost.get('currency', '')},",
            f"Cost,Max,{cost.get('max', '')},{cost.get('currency', '')},",
        ]

        return "\n".join(lines)


# Singleton
document_generator = DocumentGenerator()
