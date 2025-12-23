"""
Contract Parser Service.

Parses contract documents (PDF, DOCX) and extracts structured data:
- Work Plan (activities, milestones, deadlines)
- Budget (line items, amounts)
- Indicators (KPIs)
- Policy requirements
- Financial document templates
"""

import re
import json
from pathlib import Path
from typing import Dict, Any, List, Optional
from dataclasses import dataclass, field, asdict
from datetime import date, datetime
from enum import Enum
import hashlib


class DocumentType(str, Enum):
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    TXT = "txt"


@dataclass
class Activity:
    """Work plan activity."""
    id: str
    name: str
    description: str = ""
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    responsible: Optional[str] = None
    deliverables: List[str] = field(default_factory=list)
    status: str = "planned"  # planned, in_progress, completed, delayed


@dataclass
class Milestone:
    """Project milestone."""
    id: str
    name: str
    due_date: Optional[str] = None
    deliverables: List[str] = field(default_factory=list)
    verification: Optional[str] = None
    payment_linked: bool = False
    payment_amount: Optional[float] = None


@dataclass
class BudgetLine:
    """Budget line item."""
    id: str
    category: str
    description: str
    unit: str = "unit"
    quantity: float = 1.0
    unit_cost: float = 0.0
    total: float = 0.0
    currency: str = "USD"
    notes: Optional[str] = None


@dataclass
class Indicator:
    """Performance indicator / KPI."""
    id: str
    name: str
    description: str = ""
    baseline: Optional[float] = None
    target: Optional[float] = None
    unit: str = ""
    frequency: str = "quarterly"  # monthly, quarterly, annually
    data_source: Optional[str] = None
    responsible: Optional[str] = None


@dataclass
class PolicyRequirement:
    """Policy or compliance requirement."""
    id: str
    title: str
    description: str
    category: str  # financial, technical, reporting, compliance
    priority: str = "medium"  # critical, high, medium, low
    section_ref: Optional[str] = None


@dataclass
class DocumentTemplate:
    """Required document template from contract."""
    id: str
    name: str
    description: str = ""
    frequency: str = "once"  # once, monthly, quarterly, annually
    format: str = "pdf"
    required: bool = True


@dataclass
class ParsedContract:
    """Complete parsed contract data."""
    id: str
    filename: str
    parsed_at: str
    contract_number: Optional[str] = None
    contract_title: Optional[str] = None
    contractor: Optional[str] = None
    client: Optional[str] = None
    start_date: Optional[str] = None
    end_date: Optional[str] = None
    total_budget: Optional[float] = None
    currency: str = "USD"

    # Extracted components
    work_plan: List[Activity] = field(default_factory=list)
    milestones: List[Milestone] = field(default_factory=list)
    budget: List[BudgetLine] = field(default_factory=list)
    indicators: List[Indicator] = field(default_factory=list)
    policies: List[PolicyRequirement] = field(default_factory=list)
    document_templates: List[DocumentTemplate] = field(default_factory=list)

    # Raw extracted text sections
    raw_sections: Dict[str, str] = field(default_factory=dict)

    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary."""
        return {
            "id": self.id,
            "filename": self.filename,
            "parsed_at": self.parsed_at,
            "contract_number": self.contract_number,
            "contract_title": self.contract_title,
            "contractor": self.contractor,
            "client": self.client,
            "start_date": self.start_date,
            "end_date": self.end_date,
            "total_budget": self.total_budget,
            "currency": self.currency,
            "work_plan": [asdict(a) for a in self.work_plan],
            "milestones": [asdict(m) for m in self.milestones],
            "budget": [asdict(b) for b in self.budget],
            "indicators": [asdict(i) for i in self.indicators],
            "policies": [asdict(p) for p in self.policies],
            "document_templates": [asdict(d) for d in self.document_templates],
            "summary": {
                "activities_count": len(self.work_plan),
                "milestones_count": len(self.milestones),
                "budget_lines_count": len(self.budget),
                "indicators_count": len(self.indicators),
                "policies_count": len(self.policies),
                "templates_count": len(self.document_templates),
            }
        }


class ContractParser:
    """
    Parses contract documents and extracts structured data.

    Supports:
    - PDF parsing (via pypdf or pdfplumber)
    - DOCX parsing (via python-docx)
    - Pattern-based extraction
    - LLM-assisted extraction (optional)
    """

    # Storage for parsed contracts
    _parsed_contracts: Dict[str, ParsedContract]

    # Common patterns for extraction
    PATTERNS = {
        "contract_number": [
            r"(?:Contract|Agreement|Grant)\s*(?:No\.?|Number|#)\s*[:\s]*([A-Z0-9\-/]+)",
            r"(?:Ref|Reference)\s*[:\s]*([A-Z0-9\-/]+)",
        ],
        "date": [
            r"(\d{1,2}[./\-]\d{1,2}[./\-]\d{2,4})",
            r"(\d{4}[./\-]\d{1,2}[./\-]\d{1,2})",
            r"(\w+\s+\d{1,2},?\s+\d{4})",
        ],
        "amount": [
            r"\$\s*([\d,]+(?:\.\d{2})?)",
            r"([\d,]+(?:\.\d{2})?)\s*(?:USD|EUR|UAH)",
            r"(?:Amount|Total|Budget)[:\s]*\$?\s*([\d,]+(?:\.\d{2})?)",
        ],
        "activity": [
            r"(?:Activity|Task)\s*(\d+(?:\.\d+)?)[:\s]*(.+?)(?:\n|$)",
            r"(\d+(?:\.\d+)?)\s*[.)\-]\s*(.+?)(?:\n|$)",
        ],
        "milestone": [
            r"(?:Milestone|Deliverable)\s*(\d+)[:\s]*(.+?)(?:\n|$)",
            r"M(\d+)[:\s]*(.+?)(?:\n|$)",
        ],
        "indicator": [
            r"(?:Indicator|KPI)\s*(\d+(?:\.\d+)?)[:\s]*(.+?)(?:\n|$)",
            r"(?:Output|Outcome)\s*(\d+(?:\.\d+)?)[:\s]*(.+?)(?:\n|$)",
        ],
    }

    # Section markers for different document parts
    SECTION_MARKERS = {
        "work_plan": [
            "work plan", "workplan", "implementation plan", "activity plan",
            "scope of work", "terms of reference", "activities",
            "робочий план", "план робіт", "план діяльності"
        ],
        "budget": [
            "budget", "financial plan", "cost estimate", "pricing",
            "budget breakdown", "cost breakdown",
            "бюджет", "фінансовий план", "кошторис"
        ],
        "milestones": [
            "milestones", "deliverables", "key deliverables",
            "payment schedule", "milestone schedule",
            "етапи", "результати"
        ],
        "indicators": [
            "indicators", "kpi", "performance indicators",
            "output indicators", "outcome indicators", "results framework",
            "індикатори", "показники"
        ],
        "policies": [
            "policies", "requirements", "compliance",
            "terms and conditions", "general conditions",
            "політики", "вимоги"
        ],
        "reporting": [
            "reporting", "reports", "documentation",
            "reporting requirements", "deliverable documents",
            "звітність", "документи"
        ],
    }

    def __init__(self) -> None:
        # Instance-level storage so test runs stay isolated
        self._parsed_contracts = {}
        self._check_dependencies()

    def _check_dependencies(self):
        """Check and report available parsing libraries."""
        self.has_pypdf = False
        self.has_docx = False
        self.has_pdfplumber = False

        try:
            import pypdf
            self.has_pypdf = True
        except ImportError:
            pass

        try:
            import docx
            self.has_docx = True
        except ImportError:
            pass

        try:
            import pdfplumber
            self.has_pdfplumber = True
        except ImportError:
            pass

    def parse_file(self, file_path: str, content: Optional[bytes] = None) -> ParsedContract:
        """
        Parse a contract file and extract structured data.

        Args:
            file_path: Path to the file or filename
            content: Optional file content bytes

        Returns:
            ParsedContract with extracted data
        """
        path = Path(file_path)
        file_ext = path.suffix.lower()

        # Generate unique ID
        if content:
            file_hash = hashlib.md5(content).hexdigest()[:12]
        else:
            file_hash = hashlib.md5(path.name.encode()).hexdigest()[:12]

        contract_id = f"contract_{file_hash}"

        # Extract text based on file type
        if file_ext == ".pdf":
            text = self._extract_pdf_text(file_path, content)
        elif file_ext in (".docx", ".doc"):
            text = self._extract_docx_text(file_path, content)
        elif file_ext == ".txt":
            text = content.decode("utf-8") if content else Path(file_path).read_text()
        else:
            text = content.decode("utf-8", errors="ignore") if content else ""

        # Parse the extracted text
        parsed = self._parse_text(text, path.name, contract_id)

        # Cache the result
        self._parsed_contracts[contract_id] = parsed

        return parsed

    def _extract_pdf_text(self, file_path: str, content: Optional[bytes] = None) -> str:
        """Extract text from PDF."""
        text = ""

        if self.has_pdfplumber and content:
            import pdfplumber
            import io
            with pdfplumber.open(io.BytesIO(content)) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
                    text += "\n\n"
        elif self.has_pypdf and content:
            import pypdf
            import io
            reader = pypdf.PdfReader(io.BytesIO(content))
            for page in reader.pages:
                text += page.extract_text() or ""
                text += "\n\n"
        elif self.has_pypdf and Path(file_path).exists():
            import pypdf
            reader = pypdf.PdfReader(file_path)
            for page in reader.pages:
                text += page.extract_text() or ""
                text += "\n\n"
        else:
            text = "[PDF parsing not available - install pypdf or pdfplumber]"

        return text

    def _extract_docx_text(self, file_path: str, content: Optional[bytes] = None) -> str:
        """Extract text from DOCX."""
        text = ""

        if self.has_docx:
            import docx
            import io
            if content:
                doc = docx.Document(io.BytesIO(content))
            else:
                doc = docx.Document(file_path)

            for para in doc.paragraphs:
                text += para.text + "\n"

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
        else:
            text = "[DOCX parsing not available - install python-docx]"

        return text

    def _parse_text(self, text: str, filename: str, contract_id: str) -> ParsedContract:
        """Parse extracted text and create structured contract data."""

        parsed = ParsedContract(
            id=contract_id,
            filename=filename,
            parsed_at=datetime.now().isoformat(),
        )

        # Extract basic info
        parsed.contract_number = self._extract_contract_number(text)
        parsed.contract_title = self._extract_title(text)
        dates = self._extract_dates(text)
        if dates:
            parsed.start_date = dates[0] if len(dates) > 0 else None
            parsed.end_date = dates[1] if len(dates) > 1 else None

        amounts = self._extract_amounts(text)
        if amounts:
            parsed.total_budget = max(amounts)  # Assume largest is total

        # Split into sections
        sections = self._split_into_sections(text)
        parsed.raw_sections = sections

        # Extract structured data from sections
        if "work_plan" in sections:
            parsed.work_plan = self._extract_activities(sections["work_plan"])

        if "milestones" in sections:
            parsed.milestones = self._extract_milestones(sections["milestones"])

        if "budget" in sections:
            parsed.budget = self._extract_budget(sections["budget"])

        if "indicators" in sections:
            parsed.indicators = self._extract_indicators(sections["indicators"])

        if "policies" in sections:
            parsed.policies = self._extract_policies(sections["policies"])

        if "reporting" in sections:
            parsed.document_templates = self._extract_templates(sections["reporting"])

        # If sections not found, try extracting from full text
        if not parsed.work_plan:
            parsed.work_plan = self._extract_activities(text)
        if not parsed.budget:
            parsed.budget = self._extract_budget(text)
        if not parsed.indicators:
            parsed.indicators = self._extract_indicators(text)

        return parsed

    def _extract_contract_number(self, text: str) -> Optional[str]:
        """Extract contract number."""
        for pattern in self.PATTERNS["contract_number"]:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                return match.group(1)
        return None

    def _extract_title(self, text: str) -> Optional[str]:
        """Extract contract title from first lines."""
        lines = text.strip().split("\n")[:10]
        for line in lines:
            line = line.strip()
            if len(line) > 20 and len(line) < 200:
                # Likely a title
                if any(kw in line.lower() for kw in ["agreement", "contract", "grant", "угода", "договір"]):
                    return line
        return lines[0].strip() if lines else None

    def _extract_dates(self, text: str) -> List[str]:
        """Extract dates from text."""
        dates = []
        for pattern in self.PATTERNS["date"]:
            matches = re.findall(pattern, text)
            dates.extend(matches)
        return dates[:10]  # Limit to first 10

    def _extract_amounts(self, text: str) -> List[float]:
        """Extract monetary amounts."""
        amounts = []
        for pattern in self.PATTERNS["amount"]:
            matches = re.findall(pattern, text)
            for m in matches:
                try:
                    amount = float(m.replace(",", ""))
                    if amount > 0:
                        amounts.append(amount)
                except ValueError:
                    pass
        return sorted(set(amounts), reverse=True)

    def _split_into_sections(self, text: str) -> Dict[str, str]:
        """Split text into logical sections."""
        sections = {}
        text_lower = text.lower()

        for section_name, markers in self.SECTION_MARKERS.items():
            for marker in markers:
                # Find section start
                start_idx = text_lower.find(marker)
                if start_idx == -1:
                    continue

                # Find section end (next section or end of text)
                end_idx = len(text)
                for other_section, other_markers in self.SECTION_MARKERS.items():
                    if other_section == section_name:
                        continue
                    for other_marker in other_markers:
                        other_idx = text_lower.find(other_marker, start_idx + len(marker))
                        if other_idx != -1 and other_idx < end_idx:
                            end_idx = other_idx

                section_text = text[start_idx:end_idx].strip()
                if len(section_text) > 50:  # Only if substantial content
                    sections[section_name] = section_text
                break

        return sections

    def _extract_activities(self, text: str) -> List[Activity]:
        """Extract work plan activities."""
        activities = []

        for pattern in self.PATTERNS["activity"]:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for i, match in enumerate(matches):
                activity_id = match[0] if isinstance(match, tuple) else str(i + 1)
                name = match[1] if isinstance(match, tuple) and len(match) > 1 else str(match)

                activities.append(Activity(
                    id=f"ACT_{activity_id}".replace(".", "_"),
                    name=name.strip()[:200],
                    description="",
                ))

        # Also try line-by-line extraction
        lines = text.split("\n")
        for i, line in enumerate(lines):
            line = line.strip()
            if re.match(r"^\d+[.)]\s+", line) and len(line) > 10:
                match = re.match(r"^(\d+)[.)]\s+(.+)", line)
                if match:
                    aid = f"ACT_{match.group(1)}"
                    if not any(a.id == aid for a in activities):
                        activities.append(Activity(
                            id=aid,
                            name=match.group(2).strip()[:200],
                        ))

        return activities[:50]  # Limit

    def _extract_milestones(self, text: str) -> List[Milestone]:
        """Extract milestones."""
        milestones = []

        for pattern in self.PATTERNS["milestone"]:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                milestone_id = match[0] if isinstance(match, tuple) else "1"
                name = match[1] if isinstance(match, tuple) and len(match) > 1 else str(match)

                milestones.append(Milestone(
                    id=f"M{milestone_id}",
                    name=name.strip()[:200],
                ))

        return milestones[:20]

    def _extract_budget(self, text: str) -> List[BudgetLine]:
        """Extract budget lines."""
        budget_lines = []

        # Look for table-like structures with amounts
        lines = text.split("\n")
        for i, line in enumerate(lines):
            amounts = re.findall(r"\$?\s*([\d,]+(?:\.\d{2})?)", line)
            if amounts:
                # Try to parse as budget line
                for amount_str in amounts:
                    try:
                        amount = float(amount_str.replace(",", ""))
                        if 100 < amount < 10000000:  # Reasonable budget range
                            # Get description from start of line
                            desc = re.sub(r"\$?\s*[\d,]+(?:\.\d{2})?", "", line).strip()
                            desc = re.sub(r"[|:\-]+", " ", desc).strip()

                            if len(desc) > 5:
                                budget_lines.append(BudgetLine(
                                    id=f"BL_{len(budget_lines) + 1}",
                                    category=self._guess_category(desc),
                                    description=desc[:200],
                                    total=amount,
                                ))
                    except ValueError:
                        pass

        return budget_lines[:100]

    def _guess_category(self, description: str) -> str:
        """Guess budget category from description."""
        desc_lower = description.lower()

        if any(kw in desc_lower for kw in ["salary", "personnel", "staff", "зарплат", "персонал"]):
            return "personnel"
        elif any(kw in desc_lower for kw in ["travel", "transport", "відрядж", "транспорт"]):
            return "travel"
        elif any(kw in desc_lower for kw in ["equipment", "hardware", "software", "обладн"]):
            return "equipment"
        elif any(kw in desc_lower for kw in ["training", "workshop", "навчан"]):
            return "training"
        elif any(kw in desc_lower for kw in ["overhead", "indirect", "admin", "накладн"]):
            return "overhead"
        elif any(kw in desc_lower for kw in ["consultant", "expert", "consult", "експерт"]):
            return "consultants"
        else:
            return "other"

    def _extract_indicators(self, text: str) -> List[Indicator]:
        """Extract performance indicators."""
        indicators = []

        for pattern in self.PATTERNS["indicator"]:
            matches = re.findall(pattern, text, re.IGNORECASE | re.MULTILINE)
            for match in matches:
                ind_id = match[0] if isinstance(match, tuple) else "1"
                name = match[1] if isinstance(match, tuple) and len(match) > 1 else str(match)

                indicators.append(Indicator(
                    id=f"IND_{ind_id}".replace(".", "_"),
                    name=name.strip()[:200],
                ))

        return indicators[:50]

    def _extract_policies(self, text: str) -> List[PolicyRequirement]:
        """Extract policy requirements."""
        policies = []

        # Look for numbered requirements
        lines = text.split("\n")
        for line in lines:
            line = line.strip()
            if re.match(r"^\d+[.)]\s+", line) and len(line) > 20:
                match = re.match(r"^(\d+)[.)]\s+(.+)", line)
                if match:
                    policies.append(PolicyRequirement(
                        id=f"POL_{match.group(1)}",
                        title=match.group(2).strip()[:100],
                        description=match.group(2).strip(),
                        category=self._guess_policy_category(match.group(2)),
                    ))

        return policies[:50]

    def _guess_policy_category(self, text: str) -> str:
        """Guess policy category."""
        text_lower = text.lower()

        if any(kw in text_lower for kw in ["financial", "budget", "payment", "фінанс", "бюджет"]):
            return "financial"
        elif any(kw in text_lower for kw in ["report", "document", "submit", "звіт", "документ"]):
            return "reporting"
        elif any(kw in text_lower for kw in ["technical", "quality", "технічн", "якіст"]):
            return "technical"
        elif any(kw in text_lower for kw in ["compliance", "audit", "відповідн"]):
            return "compliance"
        else:
            return "general"

    def _extract_templates(self, text: str) -> List[DocumentTemplate]:
        """Extract required document templates."""
        templates = []

        # Common document types to look for
        doc_keywords = [
            ("progress report", "Progress Report", "quarterly"),
            ("financial report", "Financial Report", "quarterly"),
            ("final report", "Final Report", "once"),
            ("technical report", "Technical Report", "once"),
            ("audit report", "Audit Report", "annually"),
            ("work plan", "Work Plan", "once"),
            ("budget", "Budget", "once"),
            ("invoice", "Invoice", "monthly"),
            ("act of work", "Act of Work", "monthly"),
        ]

        text_lower = text.lower()
        for keyword, name, frequency in doc_keywords:
            if keyword in text_lower:
                templates.append(DocumentTemplate(
                    id=f"DOC_{name.replace(' ', '_').upper()}",
                    name=name,
                    frequency=frequency,
                    required=True,
                ))

        return templates

    def get_parsed(self, contract_id: str) -> Optional[ParsedContract]:
        """Get a previously parsed contract."""
        return self._parsed_contracts.get(contract_id)

    def list_parsed(self) -> List[Dict[str, Any]]:
        """List all parsed contracts."""
        return [
            {
                "id": c.id,
                "filename": c.filename,
                "parsed_at": c.parsed_at,
                "contract_number": c.contract_number,
                "total_budget": c.total_budget,
            }
            for c in self._parsed_contracts.values()
        ]


# Singleton instance
contract_parser = ContractParser()
