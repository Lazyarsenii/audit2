"""
Export Service - Generate PDF and Excel reports from analysis data.

Supports:
- PDF: Full analysis report with charts
- Excel: Detailed metrics and tasks spreadsheet
- Markdown: Plain text report
"""
import io
import logging
from dataclasses import dataclass
from datetime import datetime
from typing import Dict, Any, Optional, List

logger = logging.getLogger(__name__)


@dataclass
class ExportResult:
    """Result of export operation."""
    content: bytes
    filename: str
    content_type: str


class ExportService:
    """Service for exporting analysis results to various formats."""

    def __init__(self):
        self._excel_available = self._check_excel()
        self._pdf_available = self._check_pdf()
        self._word_available = self._check_word()

    def _check_word(self) -> bool:
        """Check if Word export is available."""
        try:
            from docx import Document
            return True
        except ImportError:
            logger.warning("python-docx not installed, Word export disabled")
            return False

    def _check_excel(self) -> bool:
        """Check if Excel export is available."""
        try:
            import openpyxl
            return True
        except ImportError:
            logger.warning("openpyxl not installed, Excel export disabled")
            return False

    def _check_pdf(self) -> bool:
        """Check if PDF export is available."""
        try:
            from xhtml2pdf import pisa
            return True
        except (ImportError, Exception):
            pass

        try:
            import weasyprint
            return True
        except (ImportError, OSError, Exception) as e:
            logger.warning(f"No PDF library available: {e}")
            return False

    def export_to_excel(self, analysis_data: Dict[str, Any], repo_name: str = "repo") -> ExportResult:
        """
        Export analysis results to Excel format.

        Args:
            analysis_data: Full analysis result dictionary
            repo_name: Repository name for filename

        Returns:
            ExportResult with Excel file bytes
        """
        if not self._excel_available:
            raise RuntimeError("Excel export not available - install openpyxl")

        from openpyxl import Workbook
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter

        wb = Workbook()

        # Summary Sheet
        ws_summary = wb.active
        ws_summary.title = "Summary"
        self._create_summary_sheet(ws_summary, analysis_data)

        # Metrics Sheet
        ws_metrics = wb.create_sheet("Metrics")
        self._create_metrics_sheet(ws_metrics, analysis_data)

        # Tasks Sheet
        ws_tasks = wb.create_sheet("Tasks")
        self._create_tasks_sheet(ws_tasks, analysis_data)

        # Security Sheet
        ws_security = wb.create_sheet("Security")
        self._create_security_sheet(ws_security, analysis_data)

        # Cost Estimates Sheet
        ws_cost = wb.create_sheet("Cost Estimates")
        self._create_cost_sheet(ws_cost, analysis_data)

        # Save to bytes
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{repo_name}_audit_{timestamp}.xlsx"

        return ExportResult(
            content=output.getvalue(),
            filename=filename,
            content_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )

    def _create_summary_sheet(self, ws, data: Dict[str, Any]):
        """Create summary sheet with key metrics."""
        from openpyxl.styles import Font, PatternFill, Alignment

        # Title
        ws['A1'] = "Repository Audit Report"
        ws['A1'].font = Font(size=16, bold=True)
        ws.merge_cells('A1:D1')

        # Timestamp
        ws['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        ws['A2'].font = Font(italic=True)

        # Repository info
        ws['A4'] = "Repository"
        ws['B4'] = data.get('repo_url', 'N/A')
        ws['A4'].font = Font(bold=True)

        ws['A5'] = "Analysis ID"
        ws['B5'] = data.get('analysis_id', 'N/A')
        ws['A5'].font = Font(bold=True)

        # Key Scores
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        ws['A7'] = "Category"
        ws['B7'] = "Score"
        ws['C7'] = "Max"
        ws['D7'] = "Status"

        for cell in ['A7', 'B7', 'C7', 'D7']:
            ws[cell].fill = header_fill
            ws[cell].font = header_font

        # Repo Health
        repo_health = data.get('repo_health', {})
        scores = [
            ("Documentation", repo_health.get('documentation', 0), 3),
            ("Structure", repo_health.get('structure', 0), 3),
            ("Runability", repo_health.get('runability', 0), 3),
            ("Commit History", repo_health.get('commit_history', 0), 3),
        ]

        row = 8
        for name, score, max_score in scores:
            ws[f'A{row}'] = name
            ws[f'B{row}'] = score
            ws[f'C{row}'] = max_score
            ws[f'D{row}'] = "Good" if score >= max_score * 0.7 else "Needs Work"
            row += 1

        # Total
        ws[f'A{row}'] = "Repo Health Total"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'] = repo_health.get('total', 0)
        ws[f'C{row}'] = 12

        # Tech Debt
        row += 2
        tech_debt = data.get('tech_debt', {})
        ws[f'A{row}'] = "Tech Debt Scores"
        ws[f'A{row}'].font = Font(bold=True, size=12)

        row += 1
        debt_scores = [
            ("Architecture", tech_debt.get('architecture', 0), 3),
            ("Code Quality", tech_debt.get('code_quality', 0), 3),
            ("Testing", tech_debt.get('testing', 0), 3),
            ("Infrastructure", tech_debt.get('infrastructure', 0), 3),
            ("Security", tech_debt.get('security_deps', 0), 3),
        ]

        for name, score, max_score in debt_scores:
            ws[f'A{row}'] = name
            ws[f'B{row}'] = score
            ws[f'C{row}'] = max_score
            row += 1

        ws[f'A{row}'] = "Tech Debt Total"
        ws[f'A{row}'].font = Font(bold=True)
        ws[f'B{row}'] = tech_debt.get('total', 0)
        ws[f'C{row}'] = 15

        # Product Level & Complexity
        row += 2
        ws[f'A{row}'] = "Product Level"
        ws[f'B{row}'] = data.get('product_level', 'N/A')
        ws[f'A{row}'].font = Font(bold=True)

        row += 1
        ws[f'A{row}'] = "Complexity"
        ws[f'B{row}'] = data.get('complexity', 'N/A')
        ws[f'A{row}'].font = Font(bold=True)

        # Column widths
        ws.column_dimensions['A'].width = 20
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 10
        ws.column_dimensions['D'].width = 15

    def _create_metrics_sheet(self, ws, data: Dict[str, Any]):
        """Create detailed metrics sheet."""
        from openpyxl.styles import Font, PatternFill

        ws['A1'] = "Detailed Metrics"
        ws['A1'].font = Font(size=14, bold=True)

        # Headers
        headers = ["Metric", "Value", "Category"]
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font

        # Extract metrics from various sources
        row = 4
        metrics_to_show = []

        # From cost estimate
        cost = data.get('cost_estimate', {})
        if cost:
            metrics_to_show.append(("Lines of Code", cost.get('loc', 0), "Size"))
            metrics_to_show.append(("Typical Hours", cost.get('hours_typical', 0), "Effort"))

        # From repo health
        repo_health = data.get('repo_health', {})
        for key, value in repo_health.items():
            if key != 'total':
                metrics_to_show.append((f"Repo Health: {key.replace('_', ' ').title()}", value, "Quality"))

        # From tech debt
        tech_debt = data.get('tech_debt', {})
        for key, value in tech_debt.items():
            if key != 'total':
                metrics_to_show.append((f"Tech Debt: {key.replace('_', ' ').title()}", value, "Quality"))

        for metric, value, category in metrics_to_show:
            ws.cell(row=row, column=1, value=metric)
            ws.cell(row=row, column=2, value=value)
            ws.cell(row=row, column=3, value=category)
            row += 1

        ws.column_dimensions['A'].width = 35
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15

    def _create_tasks_sheet(self, ws, data: Dict[str, Any]):
        """Create tasks backlog sheet."""
        from openpyxl.styles import Font, PatternFill

        ws['A1'] = "Improvement Tasks"
        ws['A1'].font = Font(size=14, bold=True)

        headers = ["Priority", "Category", "Title", "Description", "Estimate (hrs)"]
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        header_font = Font(color="FFFFFF", bold=True)

        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=3, column=col, value=header)
            cell.fill = header_fill
            cell.font = header_font

        tasks = data.get('tasks', [])
        row = 4

        # Priority colors
        priority_fills = {
            'P1': PatternFill(start_color="FF6B6B", end_color="FF6B6B", fill_type="solid"),
            'P2': PatternFill(start_color="FFE66D", end_color="FFE66D", fill_type="solid"),
            'P3': PatternFill(start_color="4ECDC4", end_color="4ECDC4", fill_type="solid"),
        }

        for task in tasks:
            priority = task.get('priority', 'P2')
            ws.cell(row=row, column=1, value=priority)
            if priority in priority_fills:
                ws.cell(row=row, column=1).fill = priority_fills[priority]

            ws.cell(row=row, column=2, value=task.get('category', ''))
            ws.cell(row=row, column=3, value=task.get('title', ''))
            ws.cell(row=row, column=4, value=task.get('description', '')[:200])
            ws.cell(row=row, column=5, value=task.get('estimate_hours', ''))
            row += 1

        ws.column_dimensions['A'].width = 10
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 40
        ws.column_dimensions['D'].width = 60
        ws.column_dimensions['E'].width = 15

    def _create_security_sheet(self, ws, data: Dict[str, Any]):
        """Create security findings sheet."""
        from openpyxl.styles import Font, PatternFill

        ws['A1'] = "Security Analysis"
        ws['A1'].font = Font(size=14, bold=True)

        tech_debt = data.get('tech_debt', {})
        security_score = tech_debt.get('security_deps', 3)

        ws['A3'] = "Security Score"
        ws['B3'] = f"{security_score}/3"
        ws['A3'].font = Font(bold=True)

        status_map = {0: "Critical", 1: "Warning", 2: "Good", 3: "Excellent"}
        ws['A4'] = "Status"
        ws['B4'] = status_map.get(security_score, "Unknown")

        ws['A6'] = "Security Recommendations"
        ws['A6'].font = Font(bold=True, size=12)

        recommendations = []
        if security_score < 3:
            recommendations.append("Run 'safety check' to find dependency vulnerabilities")
            recommendations.append("Run 'bandit' to find Python security issues")
            recommendations.append("Review and update outdated dependencies")
            recommendations.append("Check for hardcoded secrets and credentials")

        row = 7
        for rec in recommendations:
            ws.cell(row=row, column=1, value=f"• {rec}")
            row += 1

        ws.column_dimensions['A'].width = 60
        ws.column_dimensions['B'].width = 15

    def _create_cost_sheet(self, ws, data: Dict[str, Any]):
        """Create cost estimates sheet."""
        from openpyxl.styles import Font, PatternFill

        ws['A1'] = "Cost Estimation (COCOMO II)"
        ws['A1'].font = Font(size=14, bold=True)

        cost = data.get('cost_estimate', {})

        ws['A3'] = "Metric"
        ws['B3'] = "Value"
        ws['A3'].font = Font(bold=True)
        ws['B3'].font = Font(bold=True)

        estimates = [
            ("Lines of Code", cost.get('loc', 0)),
            ("Complexity", cost.get('complexity', 'N/A')),
            ("Hours (Minimum)", cost.get('hours_min', 0)),
            ("Hours (Typical)", cost.get('hours_typical', 0)),
            ("Hours (Maximum)", cost.get('hours_max', 0)),
        ]

        row = 4
        for label, value in estimates:
            ws.cell(row=row, column=1, value=label)
            ws.cell(row=row, column=2, value=value)
            row += 1

        # Regional estimates
        row += 1
        ws.cell(row=row, column=1, value="Regional Cost Estimates")
        ws.cell(row=row, column=1).font = Font(bold=True, size=12)

        regional = cost.get('regional_estimates', {})
        row += 1

        headers = ["Region", "Min", "Typical", "Max", "Currency"]
        header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=row, column=col, value=header)
            cell.fill = header_fill
            cell.font = Font(color="FFFFFF", bold=True)

        row += 1
        for region, values in regional.items():
            ws.cell(row=row, column=1, value=region.upper())
            ws.cell(row=row, column=2, value=values.get('min', 0))
            ws.cell(row=row, column=3, value=values.get('typical', 0))
            ws.cell(row=row, column=4, value=values.get('max', 0))
            ws.cell(row=row, column=5, value=values.get('currency', 'USD'))
            row += 1

        ws.column_dimensions['A'].width = 25
        ws.column_dimensions['B'].width = 15
        ws.column_dimensions['C'].width = 15
        ws.column_dimensions['D'].width = 15
        ws.column_dimensions['E'].width = 10

    def export_to_pdf(self, analysis_data: Dict[str, Any], repo_name: str = "repo") -> ExportResult:
        """
        Export analysis results to PDF format.

        Args:
            analysis_data: Full analysis result dictionary
            repo_name: Repository name for filename

        Returns:
            ExportResult with PDF file bytes
        """
        if not self._pdf_available:
            raise RuntimeError("PDF export not available - install xhtml2pdf or weasyprint")

        html_content = self._generate_html_report(analysis_data, repo_name)

        # Try xhtml2pdf first (pure Python, no system deps)
        try:
            from xhtml2pdf import pisa

            output = io.BytesIO()
            pisa_status = pisa.CreatePDF(html_content, dest=output)

            if pisa_status.err:
                raise RuntimeError("PDF generation failed")

            output.seek(0)

        except ImportError:
            # Fall back to weasyprint
            import weasyprint
            pdf_bytes = weasyprint.HTML(string=html_content).write_pdf()
            output = io.BytesIO(pdf_bytes or b"")

        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{repo_name}_audit_{timestamp}.pdf"

        return ExportResult(
            content=output.getvalue(),
            filename=filename,
            content_type="application/pdf"
        )

    def _generate_html_report(self, data: Dict[str, Any], repo_name: str) -> str:
        """Generate HTML content for PDF conversion."""
        repo_health = data.get('repo_health', {})
        tech_debt = data.get('tech_debt', {})
        cost = data.get('cost_estimate', {})
        tasks = data.get('tasks', [])

        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <style>
                body {{ font-family: Arial, sans-serif; margin: 40px; color: #333; }}
                h1 {{ color: #2563eb; border-bottom: 2px solid #2563eb; padding-bottom: 10px; }}
                h2 {{ color: #1e40af; margin-top: 30px; }}
                table {{ border-collapse: collapse; width: 100%; margin: 15px 0; }}
                th, td {{ border: 1px solid #ddd; padding: 10px; text-align: left; }}
                th {{ background-color: #2563eb; color: white; }}
                tr:nth-child(even) {{ background-color: #f8fafc; }}
                .score-good {{ color: #059669; font-weight: bold; }}
                .score-warn {{ color: #d97706; font-weight: bold; }}
                .score-bad {{ color: #dc2626; font-weight: bold; }}
                .summary-box {{ background: #f1f5f9; padding: 20px; border-radius: 8px; margin: 20px 0; }}
                .metric {{ display: inline-block; margin: 10px 20px; text-align: center; }}
                .metric-value {{ font-size: 24px; font-weight: bold; color: #2563eb; }}
                .metric-label {{ font-size: 12px; color: #64748b; }}
                .task-p1 {{ background-color: #fee2e2; }}
                .task-p2 {{ background-color: #fef3c7; }}
                .task-p3 {{ background-color: #d1fae5; }}
                .footer {{ margin-top: 40px; text-align: center; color: #94a3b8; font-size: 12px; }}
            </style>
        </head>
        <body>
            <h1>Repository Audit Report</h1>

            <div class="summary-box">
                <div class="metric">
                    <div class="metric-value">{repo_health.get('total', 0)}/12</div>
                    <div class="metric-label">Repo Health</div>
                </div>
                <div class="metric">
                    <div class="metric-value">{tech_debt.get('total', 0)}/15</div>
                    <div class="metric-label">Tech Debt Score</div>
                </div>
                <div class="metric">
                    <div class="metric-value">{data.get('product_level', 'N/A')}</div>
                    <div class="metric-label">Product Level</div>
                </div>
                <div class="metric">
                    <div class="metric-value">{data.get('complexity', 'N/A')}</div>
                    <div class="metric-label">Complexity</div>
                </div>
            </div>

            <p><strong>Repository:</strong> {data.get('repo_url', 'N/A')}</p>
            <p><strong>Analysis ID:</strong> {data.get('analysis_id', 'N/A')}</p>
            <p><strong>Generated:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>

            <h2>Repository Health</h2>
            <table>
                <tr><th>Category</th><th>Score</th><th>Max</th></tr>
                <tr><td>Documentation</td><td>{repo_health.get('documentation', 0)}</td><td>3</td></tr>
                <tr><td>Structure</td><td>{repo_health.get('structure', 0)}</td><td>3</td></tr>
                <tr><td>Runability</td><td>{repo_health.get('runability', 0)}</td><td>3</td></tr>
                <tr><td>Commit History</td><td>{repo_health.get('commit_history', 0)}</td><td>3</td></tr>
                <tr><th>Total</th><th>{repo_health.get('total', 0)}</th><th>12</th></tr>
            </table>

            <h2>Technical Debt</h2>
            <table>
                <tr><th>Category</th><th>Score</th><th>Max</th></tr>
                <tr><td>Architecture</td><td>{tech_debt.get('architecture', 0)}</td><td>3</td></tr>
                <tr><td>Code Quality</td><td>{tech_debt.get('code_quality', 0)}</td><td>3</td></tr>
                <tr><td>Testing</td><td>{tech_debt.get('testing', 0)}</td><td>3</td></tr>
                <tr><td>Infrastructure</td><td>{tech_debt.get('infrastructure', 0)}</td><td>3</td></tr>
                <tr><td>Security</td><td>{tech_debt.get('security_deps', 0)}</td><td>3</td></tr>
                <tr><th>Total</th><th>{tech_debt.get('total', 0)}</th><th>15</th></tr>
            </table>

            <h2>Cost Estimation</h2>
            <table>
                <tr><th>Metric</th><th>Value</th></tr>
                <tr><td>Lines of Code</td><td>{cost.get('loc', 0):,}</td></tr>
                <tr><td>Hours (Typical)</td><td>{cost.get('hours_typical', 0):.0f}</td></tr>
                <tr><td>Hours (Range)</td><td>{cost.get('hours_min', 0):.0f} - {cost.get('hours_max', 0):.0f}</td></tr>
            </table>

            <h2>Improvement Tasks ({len(tasks)} total)</h2>
            <table>
                <tr><th>Priority</th><th>Category</th><th>Task</th></tr>
                {"".join(f'<tr class="task-{t.get("priority", "p2").lower()}"><td>{t.get("priority", "")}</td><td>{t.get("category", "")}</td><td>{t.get("title", "")}</td></tr>' for t in tasks[:15])}
            </table>
            {f'<p><em>... and {len(tasks) - 15} more tasks</em></p>' if len(tasks) > 15 else ''}

            <div class="footer">
                <p>Generated by Repo Auditor | {datetime.now().strftime('%Y-%m-%d %H:%M')}</p>
            </div>
        </body>
        </html>
        """
        return html

    def export_to_markdown(self, analysis_data: Dict[str, Any], repo_name: str = "repo") -> ExportResult:
        """Export analysis results to Markdown format."""
        repo_health = analysis_data.get('repo_health', {})
        tech_debt = analysis_data.get('tech_debt', {})
        cost = analysis_data.get('cost_estimate', {})
        tasks = analysis_data.get('tasks', [])

        md = f"""# Repository Audit Report

**Repository:** {analysis_data.get('repo_url', 'N/A')}
**Generated:** {datetime.now().strftime('%Y-%m-%d %H:%M')}

## Summary

| Metric | Value |
|--------|-------|
| Repo Health | {analysis_data.get('repo_health', {}).get('total', 0)}/12 |
| Tech Debt | {tech_debt.get('total', 0)}/15 |
| Product Level | {analysis_data.get('product_level', 'N/A')} |
| Complexity | {analysis_data.get('complexity', 'N/A')} |

## Cost Estimation

- **Lines of Code:** {cost.get('loc', 0):,}
- **Hours (Typical):** {cost.get('hours_typical', 0):.0f}
- **Hours (Range):** {cost.get('hours_min', 0):.0f} - {cost.get('hours_max', 0):.0f}

## Tasks

| Priority | Category | Task |
|----------|----------|------|
"""
        for task in tasks[:20]:
            md += f"| {task.get('priority', '')} | {task.get('category', '')} | {task.get('title', '')} |\n"

        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{repo_name}_audit_{timestamp}.md"

        return ExportResult(
            content=md.encode('utf-8'),
            filename=filename,
            content_type="text/markdown"
        )

    def export_to_word(self, analysis_data: Dict[str, Any], repo_name: str = "repo") -> ExportResult:
        """
        Export analysis results to Word/DOCX format.

        Args:
            analysis_data: Full analysis result dictionary
            repo_name: Repository name for filename

        Returns:
            ExportResult with Word file bytes
        """
        if not self._word_available:
            raise RuntimeError("Word export not available - install python-docx")

        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Title
        title = doc.add_heading('Repository Audit Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle with timestamp
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.color.rgb = RGBColor(128, 128, 128)

        doc.add_paragraph()

        # Repository Info
        doc.add_heading('Repository Information', level=1)
        p = doc.add_paragraph()
        p.add_run('Repository: ').bold = True
        p.add_run(analysis_data.get('repo_url', 'N/A'))

        p = doc.add_paragraph()
        p.add_run('Analysis ID: ').bold = True
        p.add_run(analysis_data.get('analysis_id', 'N/A'))

        doc.add_paragraph()

        # Executive Summary
        doc.add_heading('Executive Summary', level=1)

        repo_health = analysis_data.get('repo_health', {})
        tech_debt = analysis_data.get('tech_debt', {})
        cost = analysis_data.get('cost_estimate', {})

        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        summary_data = [
            ("Product Level", analysis_data.get('product_level', 'N/A')),
            ("Complexity", analysis_data.get('complexity', 'N/A')),
            ("Repo Health", f"{repo_health.get('total', 0)}/12"),
            ("Tech Debt Score", f"{tech_debt.get('total', 0)}/15"),
            ("Estimated Hours", f"{cost.get('hours_typical', 0):.0f}"),
        ]

        for i, (label, value) in enumerate(summary_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)

        doc.add_paragraph()

        # Repo Health Breakdown
        doc.add_heading('Repository Health', level=1)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        health_rows = [
            ("Category", "Score", "Max"),
            ("Documentation", str(repo_health.get('documentation', 0)), "3"),
            ("Structure", str(repo_health.get('structure', 0)), "3"),
            ("Runability", str(repo_health.get('runability', 0)), "3"),
            ("Commit History", str(repo_health.get('commit_history', 0)), "3"),
        ]

        for i, row in enumerate(health_rows):
            for j, val in enumerate(row):
                table.rows[i].cells[j].text = val

        doc.add_paragraph()

        # Tech Debt Breakdown
        doc.add_heading('Technical Debt', level=1)

        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        debt_rows = [
            ("Category", "Score", "Max"),
            ("Architecture", str(tech_debt.get('architecture', 0)), "3"),
            ("Code Quality", str(tech_debt.get('code_quality', 0)), "3"),
            ("Testing", str(tech_debt.get('testing', 0)), "3"),
            ("Infrastructure", str(tech_debt.get('infrastructure', 0)), "3"),
            ("Security", str(tech_debt.get('security_deps', 0)), "3"),
        ]

        for i, row in enumerate(debt_rows):
            for j, val in enumerate(row):
                table.rows[i].cells[j].text = val

        doc.add_paragraph()

        # Cost Estimation
        doc.add_heading('Cost Estimation', level=1)

        p = doc.add_paragraph()
        p.add_run('Lines of Code: ').bold = True
        p.add_run(f"{cost.get('loc', 0):,}")

        p = doc.add_paragraph()
        p.add_run('Hours (Typical): ').bold = True
        p.add_run(f"{cost.get('hours_typical', 0):.0f}")

        p = doc.add_paragraph()
        p.add_run('Hours (Range): ').bold = True
        p.add_run(f"{cost.get('hours_min', 0):.0f} - {cost.get('hours_max', 0):.0f}")

        doc.add_paragraph()

        # Tasks
        tasks = analysis_data.get('tasks', [])
        if tasks:
            doc.add_heading(f'Improvement Tasks ({len(tasks)} total)', level=1)

            table = doc.add_table(rows=min(len(tasks) + 1, 16), cols=3)
            table.style = 'Table Grid'

            table.rows[0].cells[0].text = "Priority"
            table.rows[0].cells[1].text = "Category"
            table.rows[0].cells[2].text = "Task"

            for i, task in enumerate(tasks[:15]):
                table.rows[i + 1].cells[0].text = task.get('priority', 'P2')
                table.rows[i + 1].cells[1].text = task.get('category', '')
                table.rows[i + 1].cells[2].text = task.get('title', '')

            if len(tasks) > 15:
                p = doc.add_paragraph()
                p.add_run(f'... and {len(tasks) - 15} more tasks').italic = True

        # Footer
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Generated by Repo Auditor | {datetime.now().strftime('%Y-%m-%d %H:%M')}")
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.size = Pt(9)

        # Save to bytes
        import io
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        filename = f"{repo_name}_audit_{timestamp}.docx"

        return ExportResult(
            content=output.getvalue(),
            filename=filename,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )


# Singleton instance
export_service = ExportService()
